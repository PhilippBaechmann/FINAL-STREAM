import streamlit as st
import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document as DocxDocument
from io import BytesIO
from datetime import datetime

# Set page configuration
st.set_page_config(
    page_title="Ireland Competitor Finder", 
    page_icon="🇮🇪", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {font-size: 2.5rem; color: #1e3a8a; margin-bottom: 0.5rem;}
    .sub-header {font-size: 1.2rem; color: #475569; margin-bottom: 2rem;}
    .metric-card {background-color: #f8fafc; padding: 1rem; border-radius: 0.5rem; box-shadow: 0 1px 3px rgba(0,0,0,0.1);}
    .metric-value {font-size: 1.8rem; font-weight: bold; color: #1e40af;}
    .metric-label {font-size: 0.9rem; color: #64748b;}
    .competitor-card {background-color: #f1f5f9; padding: 1.5rem; border-radius: 0.5rem; margin-bottom: 1rem;}
    .competitor-name {font-size: 1.3rem; color: #1e3a8a; margin-bottom: 0.5rem;}
    .competitor-match {font-size: 0.9rem; color: #059669; margin-bottom: 1rem;}
    .competitor-detail {font-size: 0.95rem; margin-bottom: 0.3rem;}
</style>
""", unsafe_allow_html=True)

# Advanced text similarity function with TF-IDF weighting
@st.cache_data
def calculate_similarity_advanced(text1, text2):
    """Calculate similarity between two text strings using TF-IDF weighting"""
    import re
    from collections import Counter
    import math
    
    # Lowercase and tokenize with better tokenization
    def tokenize(text):
        # Remove non-alphanumeric
        text = re.sub(r'[^\w\s]', ' ', text.lower())
        # Split on whitespace
        tokens = text.split()
        # Remove short tokens (likely not meaningful)
        tokens = [token for token in tokens if len(token) > 2]
        # Remove common words (basic stopword removal)
        stopwords = {'and', 'the', 'for', 'with', 'that', 'this', 'our', 'are', 'not', 'has'}
        tokens = [token for token in tokens if token not in stopwords]
        return tokens
    
    tokens1 = tokenize(text1)
    tokens2 = tokenize(text2)
    
    if not tokens1 or not tokens2:
        return 0  # Handle empty cases
    
    # Count tokens
    counter1 = Counter(tokens1)
    counter2 = Counter(tokens2)
    
    # Get unique terms
    all_terms = set(counter1.keys()).union(set(counter2.keys()))
    
    # Get document frequency (for IDF)
    doc_freq = {}
    for term in all_terms:
        doc_freq[term] = (term in counter1) + (term in counter2)
    
    # Calculate dot product with TF-IDF weights
    dot_product = 0
    for term in all_terms:
        # TF-IDF weight: term frequency * inverse document frequency
        # Using log(1+tf) to dampen the effect of high frequency terms
        tf1 = math.log(1 + counter1.get(term, 0))
        tf2 = math.log(1 + counter2.get(term, 0))
        # IDF = log(total docs / docs containing term)
        idf = math.log(1 + (2 / doc_freq[term]))
        
        # Weight for each document
        weight1 = tf1 * idf
        weight2 = tf2 * idf
        
        dot_product += weight1 * weight2
    
    # Calculate magnitudes with TF-IDF weights
    magnitude1 = math.sqrt(sum((math.log(1 + counter1.get(term, 0)) * math.log(1 + (2 / doc_freq[term])))**2 for term in all_terms))
    magnitude2 = math.sqrt(sum((math.log(1 + counter2.get(term, 0)) * math.log(1 + (2 / doc_freq[term])))**2 for term in all_terms))
    
    # Calculate cosine similarity
    if magnitude1 == 0 or magnitude2 == 0:
        return 0
    else:
        return dot_product / (magnitude1 * magnitude2)

# Load and prepare data
@st.cache_data(show_spinner="Loading dataset...")
def load_data(file_path='ireland_cleaned_CHGF.xlsx'):
    try:
        df = pd.read_excel(file_path, engine='openpyxl')
        
        # Filter high-growth firms
        df['ConsistentHighGrowthFirm 2023'] = pd.to_numeric(df['ConsistentHighGrowthFirm 2023'], errors='coerce')
        high_growth_firms = df[df['ConsistentHighGrowthFirm 2023'] == 1].copy()
        
        if high_growth_firms.empty:
            st.error("No high-growth firms found in the dataset.")
            return None
            
        # Fill NAs for better processing
        high_growth_firms['Description'] = high_growth_firms['Description'].fillna('No description available')
        
        # Handle other columns
        for col in high_growth_firms.columns:
            if pd.api.types.is_numeric_dtype(high_growth_firms[col]):
                high_growth_firms[col] = high_growth_firms[col].fillna(0)
            else:
                high_growth_firms[col] = high_growth_firms[col].fillna('Not available')
                
        # Create a summary column for better matching (combines key fields)
        high_growth_firms['company_summary'] = (
            high_growth_firms['Company Name'] + ' ' + 
            high_growth_firms['Description'] + ' ' + 
            high_growth_firms['NACE_Industry'] + ' ' +
            high_growth_firms['Business_Model']
        )
        
        st.success(f"Loaded {len(high_growth_firms)} high-growth firms")
        return high_growth_firms
    except Exception as e:
        st.error(f"Error loading data: {e}")
        return None

def find_competitors(user_company_desc, df, num_results=5, industry_filter=None, business_model_filter=None):
    """Find competitors based on description similarity"""
    if df is None or len(df) == 0:
        return None, None
    
    # Apply filters if specified
    filtered_df = df.copy()
    if industry_filter and len(industry_filter) > 0:
        filtered_df = filtered_df[filtered_df['NACE_Industry'].isin(industry_filter)]
    if business_model_filter and len(business_model_filter) > 0:
        filtered_df = filtered_df[filtered_df['Business_Model'].isin(business_model_filter)]
    
    if len(filtered_df) == 0:
        st.warning("No companies match the selected filters. Showing results without filters.")
        filtered_df = df.copy()
    
    # Calculate similarity scores
    with st.spinner("Analyzing text similarity..."):
        similarities = []
        for idx, row in filtered_df.iterrows():
            # Compare user description with company summary
            company_text = row['company_summary']
            similarity = calculate_similarity_advanced(user_company_desc, company_text)
            similarities.append((idx, similarity))
        
        # Sort by similarity score (descending)
        similarities.sort(key=lambda x: x[1], reverse=True)
        
        # Get top matches
        top_indices = [idx for idx, _ in similarities[:num_results]]
        top_scores = [score for _, score in similarities[:num_results]]
        
        return filtered_df.loc[top_indices], top_scores

def create_industry_distribution(df):
    """Create industry distribution chart for competitors"""
    try:
        industry_counts = df['NACE_Industry'].value_counts().head(8)
        
        plt.figure(figsize=(10, 5))
        plt.barh(industry_counts.index, industry_counts.values, color='#3b82f6')
        plt.xlabel('Number of Companies')
        plt.title('Industry Distribution of Selected Competitors')
        plt.tight_layout()
        
        # Convert plot to image
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        return buf
    except Exception as e:
        st.error(f"Error creating chart: {e}")
        return None

def generate_docx_report(analysis_text, user_company_name, user_company_desc, competitors, similarity_scores):
    """Generate a Word document report"""
    try:
        doc = DocxDocument()
        
        # Add header
        doc.add_heading('Competitor Analysis Report', level=0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        
        # Company Information
        doc.add_heading('Your Company', level=1)
        doc.add_paragraph(f"Company Name: {user_company_name}")
        doc.add_paragraph(f"Description: {user_company_desc}")
        
        # Add competitors
        doc.add_heading('Potential Competitors', level=1)
        
        for i, ((_, row), score) in enumerate(zip(competitors.iterrows(), similarity_scores), 1):
            doc.add_heading(f"{i}. {row['Company Name']} (Match: {score*100:.1f}%)", level=2)
            
            # Company details table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Attribute'
            hdr_cells[1].text = 'Value'
            
            # Add key data to table
            key_fields = [
                'NACE_Industry', 'Business_Model', 'Estimated_Revenue_mn', 
                'Number of employees 2023', 'Growth 2023', 'City', 
                'Founded Year', 'aagr 2023', 'Public_or_Private'
            ]
            
            for key in key_fields:
                if key in row and pd.notna(row[key]) and row[key] != 'Not available':
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ')
                    row_cells[1].text = str(row[key])
            
            # Add description
            if 'Description' in row and pd.notna(row['Description']):
                doc.add_heading("Description", level=3)
                doc.add_paragraph(str(row['Description']))
            
            # Add relevance section
            doc.add_heading("Relevance", level=3)
            doc.add_paragraph(
                f"This company has a {score*100:.1f}% similarity match with your business description. "
                f"They operate in the {row['NACE_Industry']} industry"
                f"{' with a ' + row['Business_Model'] + ' business model' if row['Business_Model'] != 'Not available' else ''}."
            )
            
            doc.add_paragraph()  # Spacing
        
        # Executive summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            f"This report identifies {len(competitors)} potential competitors for {user_company_name} "
            f"from Ireland's high-growth firms database. These companies were selected based on "
            f"their similarity to your business description, industry alignment, and business model."
        )
        
        # Recommendations
        doc.add_heading('Strategic Recommendations', level=1)
        doc.add_paragraph('Based on this competitor analysis, consider the following strategic actions:')
        recommendations = [
            "Analyze these competitors' value propositions and unique selling points",
            "Identify gaps in the market that your company could address",
            "Study their business models to understand what drives growth in your sector",
            "Evaluate their digital presence, marketing strategies, and customer engagement",
            "Consider how to differentiate your offering from these established competitors",
            "Monitor their product/service developments and new market entries",
            "Explore potential partnerships or collaborations with complementary businesses"
        ]
        for rec in recommendations:
            p = doc.add_paragraph()
            p.add_run("• " + rec)
        
        # Save to BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    except Exception as e:
        st.error(f"Error generating report: {e}")
        return None

def create_radar_chart(competitors, key_metrics):
    """Create a radar chart comparing competitors on key metrics"""
    try:
        # Select metrics that are numeric
        valid_metrics = []
        for metric in key_metrics:
            if metric in competitors.columns and pd.api.types.is_numeric_dtype(competitors[metric]):
                valid_metrics.append(metric)
        
        if len(valid_metrics) < 3:
            return None  # Need at least 3 metrics for a meaningful radar chart
            
        # Normalize the data for radar chart (0-1 scale)
        radar_data = competitors[valid_metrics].copy()
        for metric in valid_metrics:
            if radar_data[metric].max() > radar_data[metric].min():
                radar_data[metric] = (radar_data[metric] - radar_data[metric].min()) / (radar_data[metric].max() - radar_data[metric].min())
            else:
                radar_data[metric] = 0.5  # If all values are the same
        
        # Create radar chart
        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(111, polar=True)
        
        # Set the angles for each metric
        angles = np.linspace(0, 2*np.pi, len(valid_metrics), endpoint=False).tolist()
        angles += angles[:1]  # Close the loop
        
        # Plot each competitor
        for i, (_, row) in enumerate(competitors.iterrows()):
            values = radar_data.loc[row.name, valid_metrics].values.flatten().tolist()
            values += values[:1]  # Close the loop
            
            ax.plot(angles, values, linewidth=2, label=row['Company Name'][:15] + '...' if len(row['Company Name']) > 15 else row['Company Name'])
            ax.fill(angles, values, alpha=0.1)
        
        # Set labels and styling
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels([m.replace('_', ' ') for m in valid_metrics])
        ax.set_yticks([0.2, 0.4, 0.6, 0.8])
        ax.set_yticklabels(['0.2', '0.4', '0.6', '0.8'])
        ax.set_title('Competitor Comparison', size=14)
        plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
        
        plt.tight_layout()
        
        # Convert plot to image
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        return buf
    except Exception as e:
        st.error(f"Error creating radar chart: {e}")
        return None

# Main application structure
def main():
    st.markdown("<h1 class='main-header'>Ireland High-Growth Competitor Finder</h1>", unsafe_allow_html=True)
    st.markdown("<p class='sub-header'>Identify potential competitors from Ireland's leading high-growth firms</p>", unsafe_allow_html=True)
    
    # Load data
    high_growth_firms = load_data()
    
    if high_growth_firms is not None:
        # Create tabs
        tab1, tab2 = st.tabs(["Competitor Analysis", "Market Insights"])
        
        with tab1:
            # Set up columns
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.subheader("Your Company Details")
                user_company_name = st.text_input("Company Name:")
                user_company_desc = st.text_area("Company Description:", height=150, 
                                                help="Describe your products, services, target market, and business model")
            
            with col2:
                st.subheader("Analysis Options")
                num_competitors = st.slider(
                    "Number of competitors to find:", 
                    min_value=3, 
                    max_value=15, 
                    value=5
                )
                
                analysis_depth = st.selectbox(
                    "Analysis detail level:",
                    options=["Standard", "Detailed", "Brief"],
                    index=0
                )
                
                st.markdown("### Optional Filters")
                
                # Industry filter with search
                industry_options = sorted(high_growth_firms['NACE_Industry'].unique())
                industry_search = st.text_input("Search industries:", placeholder="e.g. Software, Finance, etc.")
                
                if industry_search:
                    filtered_industries = [i for i in industry_options if industry_search.lower() in i.lower()]
                    industry_options = filtered_industries
                
                selected_industries = st.multiselect(
                    "Filter by Industry:",
                    options=industry_options,
                    default=None
                )
                
                # Business model filter
                if 'Business_Model' in high_growth_firms.columns:
                    bm_options = sorted(high_growth_firms['Business_Model'].unique())
                    selected_bm = st.multiselect(
                        "Filter by Business Model:",
                        options=bm_options,
                        default=None
                    )
                else:
                    selected_bm = None
            
            # Find competitors button
            if st.button("🔍 Find Competitors", type="primary"):
                if not user_company_name or not user_company_desc:
                    st.warning("Please enter both company name and description.")
                else:
                    # Find competitors
                    competitors, similarity_scores = find_competitors(
                        user_company_desc, 
                        high_growth_firms,
                        num_results=num_competitors,
                        industry_filter=selected_industries if selected_industries else None,
                        business_model_filter=selected_bm if selected_bm else None
                    )
                    
                    if competitors is not None and len(competitors) > 0:
                        # Store in session state for later use
                        st.session_state.competitors = competitors
                        st.session_state.similarity_scores = similarity_scores
                        st.session_state.user_company_name = user_company_name
                        st.session_state.user_company_desc = user_company_desc
                        
                        # Show results
                        st.markdown("## Competitor Analysis Results")
                        
                        # Summary metrics
                        metric_cols = st.columns(4)
                        with metric_cols[0]:
                            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
                            st.markdown(f"<div class='metric-value'>{len(competitors)}</div>", unsafe_allow_html=True)
                            st.markdown("<div class='metric-label'>Competitors Found</div>", unsafe_allow_html=True)
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                        with metric_cols[1]:
                            top_match = max(similarity_scores) * 100 if similarity_scores else 0
                            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
                            st.markdown(f"<div class='metric-value'>{top_match:.1f}%</div>", unsafe_allow_html=True)
                            st.markdown("<div class='metric-label'>Top Match Score</div>", unsafe_allow_html=True)
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                        with metric_cols[2]:
                            industries = competitors['NACE_Industry'].nunique()
                            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
                            st.markdown(f"<div class='metric-value'>{industries}</div>", unsafe_allow_html=True)
                            st.markdown("<div class='metric-label'>Industries</div>", unsafe_allow_html=True)
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                        with metric_cols[3]:
                            avg_growth = competitors['Growth 2023'].mean() if 'Growth 2023' in competitors.columns else 0
                            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
                            st.markdown(f"<div class='metric-value'>{avg_growth:.1f}%</div>", unsafe_allow_html=True)
                            st.markdown("<div class='metric-label'>Avg Growth (2023)</div>", unsafe_allow_html=True)
                            st.markdown("</div>", unsafe_allow_html=True)
                        
                        # Create industry distribution chart
                        industry_chart = create_industry_distribution(competitors)
                        if industry_chart:
                            st.image(industry_chart, use_column_width=True)
                        
                        # Display each competitor in a card format
                        st.markdown("### Top Matching Companies")
                        
                        for i, ((_, row), score) in enumerate(zip(competitors.iterrows(), similarity_scores)):
                            st.markdown("<div class='competitor-card'>", unsafe_allow_html=True)
                            
                            # Title row with match score
                            cols = st.columns([3, 1])
                            with cols[0]:
                                st.markdown(f"<div class='competitor-name'>{i+1}. {row['Company Name']}</div>", 
                                           unsafe_allow_html=True)
                            with cols[1]:
                                st.markdown(f"<div class='competitor-match'>Match: {score*100:.1f}%</div>", 
                                           unsafe_allow_html=True)
                            
                            # Details in columns
                            detail_cols = st.columns(2)
                            with detail_cols[0]:
                                st.markdown(f"<div class='competitor-detail'><b>Industry:</b> {row['NACE_Industry']}</div>", 
                                           unsafe_allow_html=True)
                                
                                if 'Business_Model' in row and row['Business_Model'] != 'Not available':
                                    st.markdown(f"<div class='competitor-detail'><b>Business Model:</b> {row['Business_Model']}</div>", 
                                               unsafe_allow_html=True)
                                
                                if 'City' in row and row['City'] != 'Not available':
                                    st.markdown(f"<div class='competitor-detail'><b>Location:</b> {row['City']}</div>", 
                                               unsafe_allow_html=True)
                            
                            with detail_cols[1]:
                                if 'Estimated_Revenue_mn' in row and pd.notna(row['Estimated_Revenue_mn']) and row['Estimated_Revenue_mn'] != 0:
                                    st.markdown(f"<div class='competitor-detail'><b>Est. Revenue:</b> €{row['Estimated_Revenue_mn']} million</div>", 
                                               unsafe_allow_html=True)
                                
                                if 'Growth 2023' in row and pd.notna(row['Growth 2023']) and row['Growth 2023'] != 'Not available':
                                    st.markdown(f"<div class='competitor-detail'><b>Growth (2023):</b> {row['Growth 2023']}%</div>", 
                                               unsafe_allow_html=True)
                                
                                if 'Founded Year' in row and pd.notna(row['Founded Year']) and row['Founded Year'] != 'Not available' and row['Founded Year'] != 0:
                                    st.markdown(f"<div class='competitor-detail'><b>Founded:</b> {int(row['Founded Year'])}</div>", 
                                               unsafe_allow_html=True)
                            
                            # Description - show full in detailed mode, truncated otherwise
                            if 'Description' in row and pd.notna(row['Description']):
                                desc = row['Description']
                                if analysis_depth == "Brief" and len(desc) > 150:
                                    desc = desc[:150] + "..."
                                elif analysis_depth == "Standard" and len(desc) > 300:
                                    desc = desc[:300] + "..."
                                
                                st.markdown(f"<div class='competitor-detail'><b>Description:</b> {desc}</div>", 
                                           unsafe_allow_html=True)
                            
                            # Show more/less button for detailed view 
                            if analysis_depth == "Detailed":
                                with st.expander("View Additional Details"):
                                    # Display any other interesting fields
                                    for col in competitors.columns:
                                        if col not in ['Company Name', 'NACE_Industry', 'Business_Model', 'City', 
                                                     'Estimated_Revenue_mn', 'Growth 2023', 'Founded Year', 
                                                     'Description', 'company_summary', 'ConsistentHighGrowthFirm 2023']:
                                            if pd.notna(row[col]) and row[col] != 'Not available' and row[col] != 0:
                                                st.write(f"**{col.replace('_', ' ')}:** {row[col]}")
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                        
                        # Generate and offer report download
                        report_bytes = generate_docx_report(
                            "Analysis goes here", 
                            user_company_name, 
                            user_company_desc,
                            competitors,
                            similarity_scores
                        )
                        
                        if report_bytes:
                            st.download_button(
                                label="📄 Download Detailed Report (.docx)",
                                data=report_bytes,
                                file_name=f"{user_company_name}_competitor_analysis.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    else:
                        st.error("No relevant competitors found. Try adjusting your description or filters.")
        
        with tab2:
            # Market Insights Tab
            st.subheader("Market Landscape")
            
            if 'competitors' in st.session_state:
                # Display competitors from session state
                st.write(f"Showing market insights based on {len(st.session_state.competitors)} identified competitors")
                
                # Key metrics radar chart
                st.subheader("Competitor Comparison")
                
                # Select metrics that might be interesting for comparison
                key_metrics = [
                    'Estimated_Revenue_mn', 'Growth 2023', 'aagr 2023', 
                    'Number of employees 2023', 'Founded Year'
                ]
                
                radar_chart = create_radar_chart(st.session_state.competitors, key_metrics)
                if radar_chart:
                    st.image(radar_chart, use_container_width=True)
                else:
                    st.info("Not enough numeric data available for comparison charts")
                
                # Industry breakdown
                st.subheader("Industry Breakdown")
                
                industry_counts = st.session_state.competitors['NACE_Industry'].value_counts()
                
                # Create pie chart for industries
                fig, ax = plt.subplots(figsize=(10, 6))
                industry_counts.plot(kind='pie', autopct='%1.1f%%', ax=ax)
                plt.ylabel('')
                plt.title('Industry Distribution')
                
                buf = BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                st.image(buf, use_container_width=True)
                
                # Display a simple data table
                st.subheader("Competitor Data Table")
                
                # Select relevant columns
                display_cols = [
                    'Company Name', 'NACE_Industry', 'Business_Model', 
                    'Estimated_Revenue_mn', 'Growth 2023', 'City'
                ]
                
                available_cols = [col for col in display_cols if col in st.session_state.competitors.columns]
                
                st.dataframe(
                    st.session_state.competitors[available_cols],
                    use_container_width=True
                )
                
                # Some market insights text
                st.subheader("Market Insights")
                
                # Number of companies by business model
                if 'Business_Model' in st.session_state.competitors.columns:
                    bm_counts = st.session_state.competitors['Business_Model'].value_counts()
                    if len(bm_counts) > 1:
                        most_common_bm = bm_counts.index[0]
                        st.write(f"The most common business model among competitors is **{most_common_bm}** " +
                                f"with {bm_counts[0]} out of {len(st.session_state.competitors)} companies.")
                
                # Average growth rate
                if 'Growth 2023' in st.session_state.competitors.columns:
                    avg_growth = st.session_state.competitors['Growth 2023'].mean()
                    st.write(f"The average growth rate among these competitors is **{avg_growth:.1f}%**.")
                
                # Geographic distribution
                if 'City' in st.session_state.competitors.columns:
                    city_counts = st.session_state.competitors['City'].value_counts()
                    if not city_counts.empty:
                        top_city = city_counts.index[0]
                        st.write(f"Most competitors are based in **{top_city}** " +
                                f"({city_counts[0]} out of {len(st.session_state.competitors)} companies).")
            else:
                st.info("Run a competitor analysis first to see market insights")
    else:
        st.error("Failed to load data. Please check if the Excel file exists and has the expected format.")

    # Footer
    st.markdown("---")
    st.caption("Ireland High-Growth Competitor Finder | ©️ 2025")

if __name__ == "__main__":
    main()